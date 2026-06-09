#!/usr/bin/env python3
"""
converter_md_docx.py — Timeware MPS-SW Nível C
Converte documentos .md do repositório para .docx no padrão Timeware.
Uso:
    python3 converter_md_docx.py                  # converte todos os .md em oficial/
    python3 converter_md_docx.py caminho/doc.md   # converte um arquivo específico
"""

import re, sys, os, glob
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paleta Timeware ────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x00, 0x33, 0x66)
STEEL = RGBColor(0x1A, 0x52, 0x76)
GRAY  = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ── Mapa processo ──────────────────────────────────────────────────────────────
PROCESS = {
    'ORG': 'Governança Organizacional',
    'OSW': 'Gerência Organizacional de Software',
    'GPC': 'Gerência de Processos',
    'MED': 'Medição',
    'CAP': 'Capacitação',
    'GDE': 'Gerência de Decisões',
    'GCO': 'Gerência de Configuração',
    'AQU': 'Aquisição',
    'GPR': 'Gerência de Projetos',
    'REQ': 'Engenharia de Requisitos',
    'PCP': 'Projeto e Construção do Produto',
    'ITP': 'Integração do Produto',
    'VV':  'Verificação e Validação',
}
ORDINAL = ['', '1ª', '2ª', '3ª', '4ª', '5ª']

# ── XML helpers ────────────────────────────────────────────────────────────────
def _pPr(p):
    pp = p._p.find(qn('w:pPr'))
    if pp is None:
        pp = OxmlElement('w:pPr'); p._p.insert(0, pp)
    return pp

def para_border(p, sides, color='003366', sz='6', val='single', space='1'):
    pp = _pPr(p)
    pBdr = pp.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr'); pp.append(pBdr)
    for side in sides:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), val); el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), space); el.set(qn('w:color'), color)
        pBdr.append(el)

def tab_stop(p, pos_cm, align='right'):
    pp = _pPr(p)
    tabs = pp.find(qn('w:tabs'))
    if tabs is None:
        tabs = OxmlElement('w:tabs'); pp.append(tabs)
    t = OxmlElement('w:tab')
    t.set(qn('w:val'), align); t.set(qn('w:pos'), str(int(pos_cm * 567)))
    tabs.append(t)

def keep_with_next(p):
    pp = _pPr(p)
    for tag in ('w:keepNext', 'w:keepLines'):
        pp.append(OxmlElement(tag))

def shading(cell, fill):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr'); tc.insert(0, tcPr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr'); tc.insert(0, tcPr)
    tcMar = OxmlElement('w:tcMar')
    for name, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        m = OxmlElement(f'w:{name}'); m.set(qn('w:w'), str(val)); m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

def field_code(p, code, sz=8):
    r = p.add_run()
    r.font.name = 'Calibri'; r.font.size = Pt(sz)
    fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'begin'); r._r.append(fc)
    ins = OxmlElement('w:instrText'); ins.set(qn('xml:space'), 'preserve')
    ins.text = f' {code} '; r._r.append(ins)
    fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'end'); r._r.append(fc2)

def cant_split(row):
    trPr = row._tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr'); row._tr.insert(0, trPr)
    cs = OxmlElement('w:cantSplit'); cs.set(qn('w:val'), '1'); trPr.append(cs)

def tbl_borders(table, color='AAAAAA'):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    tblB = OxmlElement('w:tblBorders')
    for name in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0');    b.set(qn('w:color'), color)
        tblB.append(b)
    tblPr.append(tblB)

# ── Inline formatter ────────────────────────────────────────────────────────────
def add_inline(p, text, sz=11, color=None):
    pattern = r'(\*\*[^*]+?\*\*|\*[^*\n]+?\*|`[^`]+?`)'
    for part in re.split(pattern, text):
        if not part: continue
        bold = italic = mono = False
        txt = part
        if part.startswith('**') and part.endswith('**'):
            txt = part[2:-2]; bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            txt = part[1:-1]; italic = True
        elif part.startswith('`') and part.endswith('`'):
            txt = part[1:-1]; mono = True
        r = p.add_run(txt)
        r.bold = bold; r.italic = italic
        r.font.name = 'Courier New' if mono else 'Calibri'
        r.font.size = Pt(10 if mono else sz)
        if color: r.font.color.rgb = color
        elif mono: r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ── Extrai metadados do .md ────────────────────────────────────────────────────
def extract_meta(md_text):
    meta = {'title': '', 'code': '', 'version': '', 'date': '', 'process': '', 'edition': '1ª Edição'}

    # Título H1
    m = re.search(r'^# (.+)', md_text, re.MULTILINE)
    if m:
        meta['title'] = m.group(1).replace('— TIMEWARE','').replace('— Timeware','').strip()

    # Primeira tabela de metadados
    for line in md_text.split('\n'):
        s = line.strip()
        if not s.startswith('|') or re.match(r'^\|[-: |]+\|$', s): continue
        cells = [c.strip() for c in s.split('|') if c.strip()]
        if len(cells) < 2: continue
        key, val = cells[0].lower(), cells[1]
        # Remove markdown bold
        val = re.sub(r'\*\*', '', val)
        if 'versão' in key or 'versao' in key:
            meta['version'] = val.strip()
            major = val.strip().split('.')[0]
            try: meta['edition'] = ORDINAL[int(major)] + ' Edição'
            except: meta['edition'] = '1ª Edição'
        elif 'data' in key:
            meta['date'] = val.strip()
        elif 'documento' in key or 'código' in key or 'codigo' in key:
            # Pega apenas o código (primeira palavra antes de ' —')
            code_raw = val.split('—')[0].strip().split()[0] if '—' in val else val.strip().split()[0]
            meta['code'] = code_raw

    # Se não extraiu code da tabela, tenta do nome do arquivo (fallback)
    return meta

def code_to_process(code):
    """PRO-GDE-001 → 'GDE' → 'Gerência de Decisões'"""
    parts = code.split('-')
    for part in parts:
        if part in PROCESS:
            return part, PROCESS[part]
    return 'ORG', 'Timeware MPS-SW'

# ── Construtor de tabela ────────────────────────────────────────────────────────
def build_table(doc, md_lines, is_meta=False):
    rows = []
    for line in md_lines:
        s = line.strip()
        if not s or re.match(r'^\|[-: |]+\|$', s): continue
        cells = [c.strip() for c in s.split('|') if c.strip() != '']
        if cells: rows.append(cells)
    if not rows: return

    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    tbl_borders(t, '003366' if not is_meta else 'AAAAAA')

    for ri, row_data in enumerate(rows):
        cant_split(t.rows[ri])
        for ci in range(ncols):
            cell = t.rows[ri].cells[ci]
            text = row_data[ci] if ci < len(row_data) else ''
            cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

            # Templates: linhas com instruções em itálico (começam com *[)
            is_instruction = text.startswith('*[') or (text.startswith('*') and text.endswith('*') and len(text) > 4)
            add_inline(p, text, sz=10)

            if ri == 0:
                for run in p.runs: run.bold = True
                if is_meta:
                    shading(cell, 'D6E4F0')
                else:
                    shading(cell, '003366')
                    for run in p.runs: run.font.color.rgb = WHITE
            elif is_meta and ci == 0:
                shading(cell, 'F0F4F8')
                for run in p.runs: run.bold = True
            elif ri % 2 == 1 and not is_meta:
                shading(cell, 'EBF5FB')
    return t

# ── Conversor principal ────────────────────────────────────────────────────────
def convert(md_path, out_path=None):
    md_text = open(md_path, encoding='utf-8').read()
    meta = extract_meta(md_text)

    # Fallback: extrai code do nome do arquivo
    if not meta['code']:
        basename = os.path.basename(md_path)
        meta['code'] = basename.split('_')[0]

    proc_key, proc_name = code_to_process(meta['code'])

    # Saída padrão: mesmo diretório, mesmo nome, extensão .docx
    if out_path is None:
        out_path = os.path.splitext(md_path)[0] + '.docx'

    doc = Document()
    sec = doc.sections[0]
    sec.page_width   = Cm(21); sec.page_height  = Cm(29.7)
    sec.top_margin   = Cm(2.8); sec.bottom_margin = Cm(2.5)
    sec.left_margin  = Cm(2.8); sec.right_margin  = Cm(2.5)
    sec.different_first_page_header_footer = True
    TW = 15.7   # largura área de texto em cm

    # Estilos globais
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].paragraph_format.space_after  = Pt(6)
    doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    doc.styles['Normal'].paragraph_format.line_spacing  = 1.15

    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'; h1.font.bold = True; h1.font.size = Pt(13)
    h1.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(18); h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'; h2.font.bold = True; h2.font.size = Pt(11)
    h2.font.color.rgb = STEEL
    h2.paragraph_format.space_before = Pt(12); h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'; h3.font.bold = True; h3.font.size = Pt(11)
    h3.font.color.rgb = STEEL; h3.font.italic = True
    h3.paragraph_format.space_before = Pt(8); h3.paragraph_format.space_after = Pt(2)
    h3.paragraph_format.keep_with_next = True

    # ── CABEÇALHO (páginas 2+) ─────────────────────────────────────────────────
    hdr = sec.header; hdr.is_linked_to_previous = False
    p = hdr.paragraphs[0]; p.clear()
    tab_stop(p, TW, 'right')
    para_border(p, ['bottom'], color='003366', sz='8', space='4')
    pp = _pPr(p)
    sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '80')
    pp.append(sp)

    r1 = p.add_run("TIMEWARE")
    r1.font.name = 'Calibri'; r1.font.size = Pt(9); r1.bold = True; r1.font.color.rgb = NAVY
    r2 = p.add_run(f"   ·   {meta['code']}")
    r2.font.name = 'Calibri'; r2.font.size = Pt(9); r2.font.color.rgb = GRAY
    p.add_run("\t")
    r3 = p.add_run(proc_name)
    r3.font.name = 'Calibri'; r3.font.size = Pt(9); r3.font.color.rgb = GRAY

    # ── CABEÇALHO página 1 (vazio) ─────────────────────────────────────────────
    fst_hdr = sec.first_page_header; fst_hdr.is_linked_to_previous = False
    fst_hdr.paragraphs[0].clear()

    # ── RODAPÉ (todas as páginas) ──────────────────────────────────────────────
    for footer in (sec.footer, sec.first_page_footer):
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]; p.clear()
        tab_stop(p, TW, 'right')
        para_border(p, ['top'], color='003366', sz='4', space='4')
        pp2 = _pPr(p)
        sp2 = OxmlElement('w:spacing'); sp2.set(qn('w:before'), '80'); sp2.set(qn('w:after'), '0')
        pp2.append(sp2)

        ver_date = f"{meta['edition']}"
        if meta['version']: ver_date = f"v{meta['version']}"
        if meta['date']:    ver_date += f"  ·  {meta['date']}"
        r_v = p.add_run(ver_date)
        r_v.font.name = 'Calibri'; r_v.font.size = Pt(8); r_v.font.color.rgb = GRAY
        p.add_run("\t")
        r_pg = p.add_run("Página ")
        r_pg.font.name = 'Calibri'; r_pg.font.size = Pt(8); r_pg.font.color.rgb = GRAY
        field_code(p, 'PAGE')
        r_de = p.add_run(" de ")
        r_de.font.name = 'Calibri'; r_de.font.size = Pt(8); r_de.font.color.rgb = GRAY
        field_code(p, 'NUMPAGES')

    # ── BLOCO TÍTULO ───────────────────────────────────────────────────────────
    bar = doc.add_paragraph()
    para_border(bar, ['top'], color='003366', sz='32', space='0')
    bar.paragraph_format.space_before = Pt(0); bar.paragraph_format.space_after = Pt(0)

    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(12); tp.paragraph_format.space_after = Pt(2)
    rt = tp.add_run(meta['title'].upper() if meta['title'] else proc_name.upper())
    rt.font.name = 'Calibri'; rt.bold = True; rt.font.size = Pt(18); rt.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.paragraph_format.space_before = Pt(0); sub.paragraph_format.space_after = Pt(2)
    parts_sub = []
    if meta['code']:    parts_sub.append(meta['code'])
    if meta['version']: parts_sub.append(f"Versão {meta['version']}")
    if meta['date']:    parts_sub.append(meta['date'])
    parts_sub.append("Timeware Brasil")
    rs = sub.add_run("   ·   ".join(parts_sub))
    rs.font.name = 'Calibri'; rs.font.size = Pt(10); rs.font.color.rgb = GRAY

    sep = doc.add_paragraph()
    para_border(sep, ['bottom'], color='003366', sz='6', space='1')
    sep.paragraph_format.space_before = Pt(6); sep.paragraph_format.space_after = Pt(14)

    # ── CONTEÚDO ───────────────────────────────────────────────────────────────
    lines = md_text.split('\n')
    table_buf = []
    first_table = True

    def flush():
        nonlocal first_table, table_buf
        if not table_buf: return
        h0 = table_buf[0] if table_buf else ''
        is_meta = first_table and 'Campo' in h0 and 'Valor' in h0
        is_hist = ('Versão' in h0 or 'Versao' in h0) and 'Data' in h0
        build_table(doc, table_buf, is_meta=(is_meta or is_hist))
        if is_meta: first_table = False
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(8)
        table_buf.clear()

    i = 0
    while i < len(lines):
        line = lines[i]

        # Ignora H1 e separadores ---
        if line.startswith('# ') or line.strip() == '---':
            flush(); i += 1; continue

        # H2 → Heading 1
        if line.startswith('## '):
            flush()
            h = doc.add_heading(line[3:].strip(), level=1)
            keep_with_next(h)
            para_border(h, ['bottom'], color='D6E4F0', sz='4', space='2')
            i += 1; continue

        # H3 → Heading 2
        if line.startswith('### '):
            flush()
            h = doc.add_heading(line[4:].strip(), level=2)
            keep_with_next(h)
            i += 1; continue

        # H4 → Heading 3
        if line.startswith('#### '):
            flush()
            h = doc.add_heading(line[5:].strip(), level=3)
            keep_with_next(h)
            i += 1; continue

        # Tabela
        if line.strip().startswith('|'):
            table_buf.append(line); i += 1
            if i >= len(lines) or not lines[i].strip().startswith('|'):
                flush()
            continue

        # Bullet list
        if re.match(r'^[-*] ', line):
            flush()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(2)
            add_inline(p, line[2:].strip())
            i += 1; continue

        # Numbered list
        if re.match(r'^\d+\. ', line):
            flush()
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(2)
            add_inline(p, re.sub(r'^\d+\. ', '', line))
            i += 1; continue

        # Blockquote > (instruções de template)
        if line.startswith('> '):
            flush()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Cm(0.8)
            para_border(p, ['left'], color='003366', sz='12', space='8')
            add_inline(p, line[2:].strip(), color=STEEL)
            i += 1; continue

        # Linha em branco
        if line.strip() == '':
            flush(); i += 1; continue

        # Parágrafo normal
        flush()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(6)
        add_inline(p, line.strip())
        i += 1

    flush()
    doc.save(out_path)
    return out_path

# ── Entry point ─────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    BASE = os.path.join(os.path.dirname(__file__), '..', 'oficial')

    if len(sys.argv) > 1:
        targets = sys.argv[1:]
    else:
        targets = glob.glob(os.path.join(BASE, '**/*.md'), recursive=True)
        # Exclui pasta de registros de projetos
        targets = [f for f in targets if '/04_registros/' not in f]

    ok = fail = 0
    for md in sorted(targets):
        try:
            out = convert(md)
            rel = os.path.relpath(out, BASE + '/..')
            print(f'  ✓ {rel}')
            ok += 1
        except Exception as e:
            print(f'  ✗ {os.path.basename(md)}: {e}')
            fail += 1

    print(f'\n{ok} convertidos, {fail} erros.')
